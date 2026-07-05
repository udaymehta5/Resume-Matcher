import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF file bytes using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")
    return "\n".join(text_content).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX file bytes using python-docx."""
    text_content = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from table cells as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")
    return "\n".join(text_content).strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatches extraction based on file extension."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore").strip()
    else:
        raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
